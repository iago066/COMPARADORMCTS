
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from collections import Counter

def export_to_docx(publicacoes, duplicados, buffer):
    doc = Document()
    titulo = doc.add_heading("Diário Consolidado - MCTS", 0)
    titulo.runs[0].font.color.rgb = RGBColor(15, 30, 68)

    doc.add_paragraph(f"Data de geração: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph(f"Total de publicações nesta parte: {len(publicacoes)}")
    doc.add_paragraph(f" - Duplicados: {len(duplicados)}")

    origens = [origem for pub in publicacoes for origem in pub.get("origens", [pub["origem"]])]
    contagem = Counter(origens)
    doc.add_paragraph("Publicações por arquivo:")
    for nome, qtd in contagem.items():
        doc.add_paragraph(f" - {nome}: {qtd}")
    doc.add_paragraph("")

    for pub in publicacoes:
        doc.add_paragraph(pub["cabecalho"], style="Intense Quote")

        doc.add_paragraph(f"Número do Processo: {pub['numero_processo']}")
        doc.add_paragraph(f"Fonte: {pub['origem']}")

        # Corpo do texto justificado linha a linha, sem excesso de espaçamento
        for linha in pub["texto"].splitlines():
            if linha.strip():
                p = doc.add_paragraph()
                run = p.add_run(linha.strip())
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        if "duplicado_de" in pub:
            doc.add_paragraph("🔁 Publicação duplicada também encontrada em:")
            for origem_info in pub["duplicado_de"]:
                doc.add_paragraph(f" - {origem_info}")

        doc.add_paragraph("")

    doc.add_paragraph("________________________________________")
    doc.add_paragraph("Legenda de cores:")
    doc.add_paragraph("Criado por Maria Clara Nogueira Diniz - OAB/PI 23765")
    doc.add_paragraph("Site: www.mcts.adv.br | E-mail: contato@mcts.adv.br")
    doc.save(buffer)
    buffer.seek(0)
